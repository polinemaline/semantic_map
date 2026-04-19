import io
import json
import mimetypes
import re
import textwrap
from datetime import datetime
from pathlib import Path
from urllib.parse import quote
from uuid import uuid4

from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from flask import (
    Blueprint,
    abort,
    current_app,
    flash,
    g,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

from .db import get_db
from .repository import (
    build_semantic_map_data,
    delete_document,
    delete_term,
    get_document,
    get_document_occurrences,
    get_stats,
    get_term,
    get_term_occurrences,
    list_conflicting_terms,
    list_documents,
    list_terms,
    save_document_with_entries,
)
from .services.document_parser import extract_text
from .services.term_extractor import extract_glossary_entries

bp = Blueprint("main", __name__)


def now_iso() -> str:
    return datetime.utcnow().isoformat(timespec="seconds")


def is_allowed_file(filename: str) -> bool:
    extension = Path(filename).suffix.lower()
    return extension in current_app.config["ALLOWED_EXTENSIONS"]


def users_count() -> int:
    db = get_db()
    row = db.execute("SELECT COUNT(*) AS count FROM users").fetchone()
    return int(row["count"])


def normalize_email(email: str) -> str:
    return email.strip().lower()


def normalize_login(login: str) -> str:
    return login.strip().lower()


def is_valid_email(email: str) -> bool:
    return re.fullmatch(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email) is not None


def get_user_by_id(user_id: int):
    db = get_db()
    return db.execute(
        """
        SELECT id, email, username, created_at
        FROM users
        WHERE id = ?
        """,
        (user_id,),
    ).fetchone()


def get_user_by_email(email: str):
    db = get_db()
    return db.execute(
        """
        SELECT id, email, username, password_hash, created_at
        FROM users
        WHERE LOWER(email) = ?
        """,
        (email,),
    ).fetchone()


def get_user_by_login(login_value: str):
    db = get_db()
    return db.execute(
        """
        SELECT id, email, username, password_hash, created_at
        FROM users
        WHERE LOWER(email) = ? OR LOWER(username) = ?
        """,
        (login_value, login_value),
    ).fetchone()


def create_user(email: str, password_hash: str) -> int:
    db = get_db()
    cursor = db.execute(
        """
        INSERT INTO users (
            username,
            display_name,
            email,
            password_hash,
            avatar_filename,
            created_at
        )
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            email,
            email.split("@")[0],
            email,
            password_hash,
            None,
            now_iso(),
        ),
    )
    db.commit()
    return int(cursor.lastrowid)


@bp.before_app_request
def load_current_user():
    user_id = session.get("user_id")
    g.current_user = get_user_by_id(user_id) if user_id else None

    auth_endpoints = {"main.login", "main.register", "static"}
    endpoint = request.endpoint or ""

    if endpoint in auth_endpoints or endpoint.startswith("static"):
        return None

    total_users = users_count()

    if total_users == 0:
        return redirect(url_for("main.register"))

    if g.current_user is None:
        return redirect(url_for("main.login"))

    return None


@bp.route("/register", methods=["GET", "POST"])
def register():
    if users_count() > 0 and g.current_user is not None:
        return redirect(url_for("main.index"))

    email_value = ""

    if request.method == "POST":
        email_value = normalize_email(request.form.get("email", ""))
        password = request.form.get("password", "")
        password_repeat = request.form.get("password_repeat", "")

        if not email_value:
            flash("Введите email.", "error")
            return render_template("register.html", email_value=email_value)

        if not is_valid_email(email_value):
            flash("Введите корректный email.", "error")
            return render_template("register.html", email_value=email_value)

        if not password:
            flash("Введите пароль.", "error")
            return render_template("register.html", email_value=email_value)

        if len(password) < 6:
            flash("Пароль должен быть не короче 6 символов.", "error")
            return render_template("register.html", email_value=email_value)

        if password != password_repeat:
            flash("Пароли не совпадают.", "error")
            return render_template("register.html", email_value=email_value)

        if get_user_by_email(email_value):
            flash("Пользователь с таким email уже существует.", "error")
            return render_template("register.html", email_value=email_value)

        user_id = create_user(
            email=email_value,
            password_hash=generate_password_hash(password),
        )

        session.clear()
        session["user_id"] = user_id

        flash("Регистрация завершена.", "success")
        return redirect(url_for("main.index"))

    return render_template("register.html", email_value=email_value)


@bp.route("/login", methods=["GET", "POST"])
def login():
    if g.current_user is not None:
        return redirect(url_for("main.index"))

    login_value = ""

    if request.method == "POST":
        login_value = normalize_login(request.form.get("login", ""))
        password = request.form.get("password", "")

        if not login_value or not password:
            flash("Введите логин и пароль.", "error")
            return render_template("login.html", login_value=login_value)

        user = get_user_by_login(login_value)

        if not user or not check_password_hash(user["password_hash"], password):
            flash("Неверный логин или пароль.", "error")
            return render_template("login.html", login_value=login_value)

        session.clear()
        session["user_id"] = int(user["id"])

        flash("Вход выполнен.", "success")
        return redirect(url_for("main.index"))

    return render_template("login.html", login_value=login_value)


@bp.route("/logout", methods=["POST"])
def logout():
    session.clear()
    return redirect(url_for("main.login"))


@bp.route("/")
def index():
    stats = get_stats()
    recent_documents = list_documents(limit=5)

    return render_template(
        "index.html",
        stats=stats,
        recent_documents=recent_documents,
    )


@bp.route("/upload", methods=["GET", "POST"])
def upload():
    if request.method == "POST":
        uploaded_file = request.files.get("file")
        title = request.form.get("title", "").strip()

        if not uploaded_file or not uploaded_file.filename:
            flash("Сначала выберите PDF или DOCX файл.", "error")
            return redirect(url_for("main.upload"))

        if not is_allowed_file(uploaded_file.filename):
            flash("Поддерживаются только PDF и DOCX.", "error")
            return redirect(url_for("main.upload"))

        original_filename = uploaded_file.filename
        safe_name = secure_filename(original_filename)
        extension = Path(safe_name).suffix.lower()
        saved_filename = f"{uuid4().hex}{extension}"
        save_path = Path(current_app.config["UPLOAD_FOLDER"]) / saved_filename

        try:
            uploaded_file.save(save_path)

            text = extract_text(save_path)

            if not text.strip():
                raise ValueError(
                    "Не удалось извлечь текст. Для MVP нужны текстовые PDF или DOCX."
                )

            entries = extract_glossary_entries(text)

            document_id = save_document_with_entries(
                title=title or Path(original_filename).stem,
                original_filename=original_filename,
                saved_filename=saved_filename,
                file_type=extension.lstrip("."),
                full_text=text,
                entries=entries,
            )

            flash(
                f"Документ загружен. Найдено терминов: {len(entries)}.",
                "success",
            )

            return redirect(url_for("main.document_detail", document_id=document_id))

        except Exception as exc:
            if save_path.exists():
                save_path.unlink()

            flash(f"Ошибка обработки файла: {exc}", "error")
            return redirect(url_for("main.upload"))

    return render_template("upload.html")


@bp.route("/documents")
def documents():
    documents_list = list_documents()
    return render_template("documents.html", documents=documents_list)


@bp.route("/documents/<int:document_id>")
def document_detail(document_id: int):
    document = get_document(document_id)

    if not document:
        abort(404)

    occurrences = get_document_occurrences(document_id)

    return render_template(
        "document_detail.html",
        document=document,
        occurrences=occurrences,
    )


@bp.route("/documents/<int:document_id>/source")
def document_source(document_id: int):
    document = get_document(document_id)

    if not document:
        abort(404)

    return render_template("source_view.html", document=document)


@bp.route("/documents/<int:document_id>/source-file")
def document_source_file(document_id: int):
    document = get_document(document_id)

    if not document:
        abort(404)

    file_path = Path(current_app.config["UPLOAD_FOLDER"]) / document["saved_filename"]

    if not file_path.exists():
        abort(404)

    mimetype, _ = mimetypes.guess_type(document["original_filename"])

    response = send_file(
        file_path,
        mimetype=mimetype,
        as_attachment=False,
        download_name=document["original_filename"],
        max_age=0,
    )

    quoted_filename = quote(document["original_filename"])
    response.headers["Content-Disposition"] = (
        f"inline; filename*=UTF-8''{quoted_filename}"
    )

    return response


@bp.route("/documents/<int:document_id>/delete", methods=["POST"])
def document_delete(document_id: int):
    deleted = delete_document(document_id)

    if not deleted:
        abort(404)

    file_path = Path(current_app.config["UPLOAD_FOLDER"]) / deleted["saved_filename"]

    if file_path.exists():
        file_path.unlink()

    flash(f"Документ «{deleted['title']}» удалён.", "success")
    return redirect(url_for("main.documents"))


@bp.route("/documents/<int:document_id>/report.docx")
def document_report_docx(document_id: int):
    document = get_document(document_id)

    if not document:
        abort(404)

    occurrences = get_document_occurrences(document_id)
    conflict_groups = build_document_conflict_groups(occurrences)
    map_data = build_semantic_map_data(document_id=document_id)

    report_file = create_document_report_docx(
        document=document,
        occurrences=occurrences,
        conflict_groups=conflict_groups,
        map_data=map_data,
    )

    safe_title = secure_filename(document["title"]) or f"document_{document_id}"
    download_name = f"{safe_title}_report.docx"

    return send_file(
        report_file,
        as_attachment=True,
        download_name=download_name,
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )


@bp.route("/terms")
def terms():
    query = request.args.get("q", "").strip()
    terms_list = list_terms(query)

    return render_template(
        "terms.html",
        terms=terms_list,
        query=query,
    )


@bp.route("/terms/<int:term_id>")
def term_detail(term_id: int):
    term = get_term(term_id)

    if not term:
        abort(404)

    occurrences = get_term_occurrences(term_id)
    grouped_definitions: list[dict] = []
    index_by_normalized: dict[str, dict] = {}

    for row in occurrences:
        normalized_key = row["definition_normalized"]

        if normalized_key not in index_by_normalized:
            group = {
                "definition_text": row["definition_text"],
                "documents": [],
                "source_lines": [],
            }
            index_by_normalized[normalized_key] = group
            grouped_definitions.append(group)

        index_by_normalized[normalized_key]["documents"].append(
            {
                "id": row["document_id"],
                "title": row["document_title"],
                "filename": row["original_filename"],
            }
        )
        index_by_normalized[normalized_key]["source_lines"].append(row["source_line"])

    has_conflict = len(grouped_definitions) > 1

    return render_template(
        "term_detail.html",
        term=term,
        grouped_definitions=grouped_definitions,
        has_conflict=has_conflict,
    )


@bp.route("/terms/<int:term_id>/delete", methods=["POST"])
def term_delete(term_id: int):
    deleted = delete_term(term_id)

    if not deleted:
        abort(404)

    flash(f"Термин «{deleted['name']}» удалён.", "success")
    return redirect(url_for("main.terms"))


@bp.route("/compare")
def compare():
    conflicts = list_conflicting_terms()
    return render_template("compare.html", conflicts=conflicts)


@bp.route("/semantic-map")
def semantic_map():
    selected_term_id = request.args.get("term_id", type=int)
    selected_document_id = request.args.get("document_id", type=int)

    if selected_term_id:
        selected_document_id = None

    if selected_term_id is not None:
        term = get_term(selected_term_id)

        if not term:
            abort(404)

    if selected_document_id is not None:
        document = get_document(selected_document_id)

        if not document:
            abort(404)

    map_data = build_semantic_map_data(
        term_id=selected_term_id,
        document_id=selected_document_id,
    )

    return render_template(
        "semantic_map.html",
        documents=list_documents(),
        terms=list_terms(),
        selected_term_id=selected_term_id,
        selected_document_id=selected_document_id,
        map_summary=map_data["summary"],
        nodes_json=json.dumps(map_data["nodes"], ensure_ascii=False),
        edges_json=json.dumps(map_data["edges"], ensure_ascii=False),
    )


def build_document_conflict_groups(occurrences) -> list[dict]:
    term_ids = sorted({int(row["term_id"]) for row in occurrences})
    groups: list[dict] = []

    for term_id in term_ids:
        term = get_term(term_id)

        if not term or int(term["definitions_count"]) <= 1:
            continue

        term_occurrences = get_term_occurrences(term_id)
        definition_map: dict[str, dict] = {}

        for row in term_occurrences:
            normalized = row["definition_normalized"]

            if normalized not in definition_map:
                definition_map[normalized] = {
                    "definition_text": row["definition_text"],
                    "documents": [],
                    "_document_ids": set(),
                }

            document_id = int(row["document_id"])

            if document_id not in definition_map[normalized]["_document_ids"]:
                definition_map[normalized]["_document_ids"].add(document_id)
                definition_map[normalized]["documents"].append(
                    {
                        "id": document_id,
                        "title": row["document_title"],
                        "filename": row["original_filename"],
                    }
                )

        definitions = list(definition_map.values())

        for definition in definitions:
            definition.pop("_document_ids", None)
            definition["documents"].sort(key=lambda item: item["title"].lower())

        definitions.sort(key=lambda item: item["definition_text"].lower())

        groups.append(
            {
                "term_id": term_id,
                "term_name": term["name"],
                "definitions_count": len(definitions),
                "documents_count": int(term["documents_count"]),
                "definitions": definitions,
            }
        )

    groups.sort(key=lambda item: item["term_name"].lower())
    return groups


def create_document_report_docx(
    *,
    document,
    occurrences,
    conflict_groups: list[dict],
    map_data: dict,
) -> io.BytesIO:
    docx = DocxDocument()

    set_docx_base_styles(docx)

    title = docx.add_heading("Отчет по документу", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    docx.add_paragraph(f"Документ: {document['title']}")
    docx.add_paragraph(f"Исходный файл: {document['original_filename']}")
    docx.add_paragraph(f"Формат: {document['file_type'].upper()}")
    docx.add_paragraph(f"Дата загрузки: {document['created_at']}")

    total_terms = len({int(row["term_id"]) for row in occurrences})
    total_definitions = len(occurrences)
    conflict_count = len(conflict_groups)

    docx.add_heading("Сводные показатели", level=1)

    stats_table = docx.add_table(rows=1, cols=2)
    stats_table.style = "Table Grid"
    stats_table.rows[0].cells[0].text = "Общее количество терминов"
    stats_table.rows[0].cells[1].text = str(total_terms)

    add_docx_table_row(stats_table, "Количество определений", str(total_definitions))
    add_docx_table_row(
        stats_table,
        "Количество конфликтующих терминов",
        str(conflict_count),
    )

    docx.add_heading("Все определения документа", level=1)

    if occurrences:
        definitions_table = docx.add_table(rows=1, cols=2)
        definitions_table.style = "Table Grid"
        definitions_table.rows[0].cells[0].text = "Термин"
        definitions_table.rows[0].cells[1].text = "Определение"

        for row in occurrences:
            add_docx_table_row(
                definitions_table,
                row["term_name"],
                row["definition_text"],
            )
    else:
        docx.add_paragraph("В документе не найдено извлеченных определений.")

    docx.add_heading("Конфликтующие термины", level=1)

    if conflict_groups:
        for group in conflict_groups:
            docx.add_heading(group["term_name"], level=2)
            docx.add_paragraph(
                f"Разных определений: {group['definitions_count']}. "
                f"Документов с термином: {group['documents_count']}."
            )

            for index, definition in enumerate(group["definitions"], start=1):
                docx.add_paragraph(f"Определение {index}", style="List Bullet")
                docx.add_paragraph(definition["definition_text"])

                docs_label = ", ".join(doc["title"] for doc in definition["documents"])
                docx.add_paragraph(f"Документы: {docs_label}")
    else:
        docx.add_paragraph(
            "Для терминов из этого документа конфликтующие определения не обнаружены."
        )

    map_section = docx.add_section(WD_SECTION.NEW_PAGE)
    map_section.orientation = WD_ORIENT.LANDSCAPE

    if map_section.page_width < map_section.page_height:
        map_section.page_width, map_section.page_height = (
            map_section.page_height,
            map_section.page_width,
        )

    map_section.left_margin = Inches(0.35)
    map_section.right_margin = Inches(0.35)
    map_section.top_margin = Inches(0.35)
    map_section.bottom_margin = Inches(0.35)

    heading = docx.add_heading("Семантическая карта документа", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        map_images = create_semantic_map_png_pages(map_data, rows_per_page=5)
        available_width = (
            map_section.page_width
            - map_section.left_margin
            - map_section.right_margin
        )

        for index, map_image in enumerate(map_images, start=1):
            if index > 1:
                docx.add_page_break()

            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(map_image, width=available_width)

    except Exception as exc:
        docx.add_paragraph(f"Не удалось добавить изображение карты: {exc}")

    output = io.BytesIO()
    docx.save(output)
    output.seek(0)

    return output


def set_docx_base_styles(docx: DocxDocument) -> None:
    normal_style = docx.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = docx.styles[style_name]
        style.font.name = "Arial"


def add_docx_table_row(table, left: str, right: str) -> None:
    row = table.add_row()
    row.cells[0].text = str(left)
    row.cells[1].text = str(right)


def create_semantic_map_png_pages(
    map_data: dict,
    rows_per_page: int = 5,
) -> list[io.BytesIO]:
    nodes = map_data.get("nodes", [])
    edges = map_data.get("edges", [])

    node_by_id = {node["id"]: node for node in nodes}
    document_nodes = [node for node in nodes if node["type"] == "document"]

    term_nodes = sorted(
        [node for node in nodes if node["type"] == "term"],
        key=lambda item: item["label"].lower(),
    )

    term_to_definition: dict[str, str] = {}

    for edge in edges:
        if edge["label"] == "имеет определение":
            term_to_definition[edge["source"]] = edge["target"]

    rows = []

    for term_node in term_nodes:
        definition_id = term_to_definition.get(term_node["id"])
        definition_node = node_by_id.get(definition_id) if definition_id else None

        rows.append(
            {
                "term": term_node,
                "definition": definition_node,
            }
        )

    if not rows:
        return [
            create_semantic_map_png_page(
                document_nodes=document_nodes,
                rows=[],
                edges=edges,
                page_number=1,
                pages_count=1,
            )
        ]

    pages = []
    pages_count = (len(rows) + rows_per_page - 1) // rows_per_page

    for page_index in range(pages_count):
        start = page_index * rows_per_page
        end = start + rows_per_page

        pages.append(
            create_semantic_map_png_page(
                document_nodes=document_nodes,
                rows=rows[start:end],
                edges=edges,
                page_number=page_index + 1,
                pages_count=pages_count,
            )
        )

    return pages


def create_semantic_map_png_page(
    *,
    document_nodes: list[dict],
    rows: list[dict],
    edges: list[dict],
    page_number: int,
    pages_count: int,
) -> io.BytesIO:
    from PIL import Image, ImageDraw, ImageFont

    font_regular = load_report_font(ImageFont, 26)
    font_bold = load_report_font(ImageFont, 28)
    font_small = load_report_font(ImageFont, 20)

    card_widths = {
        "document": 360,
        "term": 420,
        "definition": 980,
    }

    x_document = 90
    x_term = 620
    x_definition = 1120

    row_gap = 54
    top = 120

    measure_image = Image.new("RGB", (10, 10), "#ffffff")
    measure_draw = ImageDraw.Draw(measure_image)

    measured_rows = []

    for row in rows:
        term_height = estimate_card_height_by_pixels(
            measure_draw,
            row["term"]["label"],
            font_bold,
            card_widths["term"] - 56,
            118,
        )

        if row["definition"]:
            definition_height = estimate_card_height_by_pixels(
                measure_draw,
                row["definition"]["label"],
                font_regular,
                card_widths["definition"] - 64,
                148,
            )
        else:
            definition_height = 148

        row_height = max(term_height, definition_height)

        measured_rows.append(
            {
                **row,
                "term_height": term_height,
                "definition_height": definition_height,
                "row_height": row_height,
            }
        )

    if measured_rows:
        content_height = sum(row["row_height"] for row in measured_rows)
        content_height += max(0, len(measured_rows) - 1) * row_gap
    else:
        content_height = 220

    width = 2200
    height = max(1150, top + content_height + 180)

    image = Image.new("RGB", (width, height), "#f6f9ff")
    draw = ImageDraw.Draw(image)

    draw_report_grid(draw, width, height)

    positions: dict[str, dict] = {}

    document_label = document_nodes[0]["label"] if document_nodes else "Документ"
    document_height = estimate_card_height_by_pixels(
        draw,
        document_label,
        font_bold,
        card_widths["document"] - 56,
        118,
    )
    document_y = top + max(content_height / 2 - document_height / 2, 0)

    if document_nodes:
        positions[document_nodes[0]["id"]] = {
            "x": x_document,
            "y": int(document_y),
            "width": card_widths["document"],
            "height": document_height,
        }

    cursor_y = top

    for row in measured_rows:
        row_center = cursor_y + row["row_height"] / 2

        term_y = int(row_center - row["term_height"] / 2)
        positions[row["term"]["id"]] = {
            "x": x_term,
            "y": term_y,
            "width": card_widths["term"],
            "height": row["term_height"],
        }

        if row["definition"]:
            definition_y = int(row_center - row["definition_height"] / 2)
            positions[row["definition"]["id"]] = {
                "x": x_definition,
                "y": definition_y,
                "width": card_widths["definition"],
                "height": row["definition_height"],
            }

        cursor_y += row["row_height"] + row_gap

    for edge in edges:
        if edge["label"] == "упоминает термин":
            continue

        source = positions.get(edge["source"])
        target = positions.get(edge["target"])

        if not source or not target:
            continue

        color = "#9aa8bd" if edge["label"] == "содержит термин" else "#8057d8"
        draw_report_connection(draw, source, target, color=color, dashed=False)

    for edge in edges:
        if edge["label"] != "упоминает термин":
            continue

        source = positions.get(edge["source"])
        target = positions.get(edge["target"])

        if not source or not target:
            continue

        draw_report_connection(draw, source, target, color="#111827", dashed=True)

    if document_nodes:
        draw_report_card(
            draw,
            positions[document_nodes[0]["id"]],
            document_label,
            fill="#ffffff",
            outline="#3d72f6",
            accent="#3d72f6",
            font=font_bold,
            wrap_width_px=card_widths["document"] - 56,
        )

    for row in measured_rows:
        term = row["term"]
        term_color = get_report_term_color(term)

        draw_report_card(
            draw,
            positions[term["id"]],
            term["label"],
            fill="#ffffff",
            outline=term_color,
            accent=term_color,
            font=font_bold,
            wrap_width_px=card_widths["term"] - 56,
        )

        definition = row["definition"]

        if definition:
            draw_report_card(
                draw,
                positions[definition["id"]],
                definition["label"],
                fill="#ffffff",
                outline="#8057d8",
                accent="#8057d8",
                font=font_regular,
                wrap_width_px=card_widths["definition"] - 64,
            )

    if pages_count > 1:
        page_label = f"Фрагмент карты {page_number} из {pages_count}"
        draw.text(
            (90, height - 70),
            page_label,
            fill="#5d6b82",
            font=font_small,
        )

    output = io.BytesIO()
    image.save(output, format="PNG", dpi=(220, 220))
    output.seek(0)

    return output


def load_report_font(ImageFont, size: int):
    candidates = [
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]

    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            continue

    return ImageFont.load_default()


def wrap_text_to_pixels(draw, text: str, font, max_width: int) -> list[str]:
    words = str(text).split()

    if not words:
        return [""]

    lines: list[str] = []
    current = ""

    for word in words:
        candidate = f"{current} {word}".strip()

        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
            continue

        if current:
            lines.append(current)
            current = ""

        if draw.textlength(word, font=font) <= max_width:
            current = word
            continue

        chunk = ""

        for char in word:
            candidate_chunk = f"{chunk}{char}"

            if draw.textlength(candidate_chunk, font=font) <= max_width:
                chunk = candidate_chunk
            else:
                if chunk:
                    lines.append(chunk)
                chunk = char

        current = chunk

    if current:
        lines.append(current)

    return lines


def get_font_line_height(draw, font) -> int:
    bbox = draw.textbbox((0, 0), "АБВabc123", font=font)
    return bbox[3] - bbox[1]


def estimate_card_height_by_pixels(
    draw,
    text: str,
    font,
    max_width: int,
    min_height: int,
) -> int:
    lines = wrap_text_to_pixels(draw, text, font, max_width)
    line_height = get_font_line_height(draw, font) + 10
    return max(min_height, 74 + len(lines) * line_height)


def draw_report_grid(draw, width: int, height: int) -> None:
    step = 32
    color = "#dfe7f3"

    for x in range(0, width, step):
        draw.line((x, 0, x, height), fill=color, width=1)

    for y in range(0, height, step):
        draw.line((0, y, width, y), fill=color, width=1)


def draw_report_card(
    draw,
    pos: dict,
    text: str,
    *,
    fill: str,
    outline: str,
    accent: str,
    font,
    wrap_width_px: int,
) -> None:
    x = pos["x"]
    y = pos["y"]
    width = pos["width"]
    height = pos["height"]

    radius = 18

    draw.rounded_rectangle(
        (x, y, x + width, y + height),
        radius=radius,
        fill=fill,
        outline=outline,
        width=2,
    )

    draw.rounded_rectangle(
        (x + 16, y + 14, x + width - 16, y + 18),
        radius=999,
        fill=accent,
    )

    lines = wrap_text_to_pixels(draw, str(text), font, wrap_width_px)
    line_height = get_font_line_height(draw, font) + 10
    total_text_height = len(lines) * line_height
    text_y = y + height / 2 - total_text_height / 2 + 4

    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        text_width = bbox[2] - bbox[0]

        draw.text(
            (x + width / 2 - text_width / 2, text_y),
            line,
            fill="#152033",
            font=font,
        )

        text_y += line_height


def draw_report_connection(
    draw,
    source: dict,
    target: dict,
    *,
    color: str,
    dashed: bool,
) -> None:
    if source["x"] <= target["x"]:
        start = (
            source["x"] + source["width"],
            source["y"] + source["height"] / 2,
        )
        end = (
            target["x"],
            target["y"] + target["height"] / 2,
        )
    else:
        start = (
            source["x"],
            source["y"] + source["height"] / 2,
        )
        end = (
            target["x"] + target["width"],
            target["y"] + target["height"] / 2,
        )

    mid_x = start[0] + (end[0] - start[0]) / 2

    points = [
        start,
        (mid_x, start[1]),
        (mid_x, end[1]),
        end,
    ]

    if dashed:
        draw_dashed_polyline(draw, points, fill=color, width=3)
    else:
        draw.line(points, fill=color, width=3)


def draw_dashed_polyline(draw, points, *, fill: str, width: int) -> None:
    for start, end in zip(points, points[1:]):
        draw_dashed_segment(draw, start, end, fill=fill, width=width)


def draw_dashed_segment(draw, start, end, *, fill: str, width: int) -> None:
    x1, y1 = start
    x2, y2 = end

    dash = 10
    gap = 7

    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 >= x1 else -1
        x = x1

        while (direction == 1 and x < x2) or (direction == -1 and x > x2):
            dash_end = x + direction * dash

            if direction == 1:
                dash_end = min(dash_end, x2)
            else:
                dash_end = max(dash_end, x2)

            draw.line((x, y1, dash_end, y2), fill=fill, width=width)
            x = dash_end + direction * gap
    else:
        direction = 1 if y2 >= y1 else -1
        y = y1

        while (direction == 1 and y < y2) or (direction == -1 and y > y2):
            dash_end = y + direction * dash

            if direction == 1:
                dash_end = min(dash_end, y2)
            else:
                dash_end = max(dash_end, y2)

            draw.line((x1, y, x2, dash_end), fill=fill, width=width)
            y = dash_end + direction * gap


def get_report_term_color(term_node: dict) -> str:
    raw_color = str(term_node.get("color", "")).lower()

    if "dc2626" in raw_color or "red" in raw_color:
        return "#df4565"

    if "16a34a" in raw_color or "green" in raw_color:
        return "#169b68"

    return "#646bff"