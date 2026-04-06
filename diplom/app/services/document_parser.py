from pathlib import Path

import pymupdf
from docx import Document as DocxDocument


class UnsupportedFileTypeError(Exception):
    pass


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(path)

    if suffix == ".docx":
        return extract_text_from_docx(path)

    raise UnsupportedFileTypeError(f"Неподдерживаемый формат файла: {suffix}")


def extract_text_from_pdf(file_path: Path) -> str:
    document = pymupdf.open(file_path)
    pages: list[str] = []

    try:
        for page in document:
            pages.append(page.get_text("text", sort=True))
    finally:
        document.close()

    return "\n".join(pages).strip()


def extract_text_from_docx(file_path: Path) -> str:
    document = DocxDocument(file_path)
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks).strip()